from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.section import Section
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


CHARS_TO_REMOVE = ("#", "-","*")


def clean_text(text: str) -> str:
    for char in CHARS_TO_REMOVE:
        text = text.replace(char, "")
    return text


def clean_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = clean_text(run.text)


def clean_table(table: Table) -> None:
    for row in table.rows:
        for cell in row.cells:
            clean_container(cell)


def clean_container(container: _Document | _Cell) -> None:
    for paragraph in container.paragraphs:
        clean_paragraph(paragraph)
    for table in container.tables:
        clean_table(table)


def clean_headers_and_footers(section: Section) -> None:
    clean_container(section.header)
    clean_container(section.footer)


def clean_docx(input_path: Path, output_path: Path) -> None:
    doc = Document(str(input_path))
    clean_container(doc)
    for section in doc.sections:
        clean_headers_and_footers(section)
    doc.save(str(output_path))


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="去除 DOCX 文本中的 # 和 - 字符。",
    )
    parser.add_argument("input", type=Path, help="输入 DOCX 文件路径")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        help="输出 DOCX 文件路径（默认：原文件名_cleaned.docx）",
    )
    return parser


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()

    input_path: Path = args.input
    if input_path.suffix.lower() != ".docx":
        raise ValueError("输入文件必须是 .docx 格式")
    if not input_path.exists():
        raise FileNotFoundError(f"找不到输入文件: {input_path}")

    output_path: Path = (
        args.output if args.output else input_path.with_name(f"{input_path.stem}_cleaned.docx")
    )
    clean_docx(input_path, output_path)
    print(f"处理完成，输出文件: {output_path}")


if __name__ == "__main__":
    main()
